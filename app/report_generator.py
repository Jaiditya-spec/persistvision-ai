from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.gemini_responder import generate_assumption_insights, generate_block_commentaries

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "outputs"

DURATION_LABELS = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10+"]
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x10, 0x24, 0x3E)


def _set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_font(run, size=9, bold=False, italic=False, color=None, name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    if color:
        run.font.color.rgb = color


def generate_word_report(cohort_results, full_breakdown):
    """
    Builds the Assumption Setting Word report: a title page and
    Gemini-written executive summary in portrait, followed by the full
    duration-banded working for every ERA x Channel block in landscape,
    each with a Gemini-written commentary paragraph below it.
    """

    OUTPUT_DIR.mkdir(exist_ok=True)

    insights_text = generate_assumption_insights(cohort_results)
    if not insights_text:
        insights_text = (
            "Executive summary could not be generated automatically for this run. "
            "See the detailed tables below and the accompanying Excel summary for "
            "the full cohort-level results."
        )

    commentary_map = generate_block_commentaries(full_breakdown)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    doc.add_heading("Assumption Setting Report", level=0)
    subtitle = doc.add_paragraph("PersistVision AI — Actuarial Persistency Intelligence")
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(0x4A, 0x5D, 0x78)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph("Periods analysed: YTD Oct'25 vs YTD Jun'26")

    doc.add_heading("Executive Summary", level=1)
    for paragraph_text in insights_text.split("\n"):
        if paragraph_text.strip():
            doc.add_paragraph(paragraph_text.strip())

    doc.add_section()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width, section.page_height = new_w, new_h
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    doc.add_heading("Duration-Banded Assumption Detail", level=1)
    doc.add_paragraph(
        "The following tables show the full assumption-setting calculation for "
        "each ERA x Channel combination, broken out by Short Pay and Long Pay, "
        "across all policy durations, with commentary on notable movements."
    )

    usable_width = section.page_width - section.left_margin - section.right_margin
    label_w = Inches(0.75)
    metric_w = Inches(1.7)
    dur_w = int((usable_width - label_w - metric_w) / 10)

    for block in full_breakdown:
        era = block["era"]
        channel = block["channel"]

        doc.add_heading(f"{era} — {channel} Channel", level=2)

        table = doc.add_table(rows=1, cols=12)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        widths = [label_w, metric_w] + [dur_w] * 10

        hdr = table.rows[0].cells
        hdr[0].text = "Pay Type"
        hdr[1].text = "Metric"
        for i, d in enumerate(DURATION_LABELS):
            hdr[2 + i].text = d

        for i, cell in enumerate(hdr):
            cell.width = widths[i]
            _set_cell_shading(cell, "10243E")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if not p.runs:
                    p.add_run("")
                for r in p.runs:
                    _set_font(r, size=9, bold=True, color=WHITE)

        for pay_type, metrics in block["pay_data"].items():
            first_row = True
            for metric_name, values in metrics.items():
                row_cells = table.add_row().cells
                row_cells[0].text = pay_type if first_row else ""
                row_cells[1].text = metric_name
                is_proposed = metric_name == "Proposed Assumption"

                for i, val in enumerate(values):
                    row_cells[2 + i].text = f"{val*100:.1f}%" if val is not None else "N/A"

                for i, cell in enumerate(row_cells):
                    cell.width = widths[i]
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if not p.runs:
                            p.add_run("")
                        for r in p.runs:
                            _set_font(r, size=9, bold=is_proposed)
                    if is_proposed:
                        _set_cell_shading(cell, "E3F3EA")

                first_row = False

        commentary = commentary_map.get((era, channel))
        if not commentary:
            commentary = (
                "Commentary unavailable for this block — review the figures above "
                "directly against the prior assumption before adoption."
            )

        label_p = doc.add_paragraph()
        label_run = label_p.add_run("Commentary: ")
        _set_font(label_run, size=9, bold=True, color=INK)

        comment_run = label_p.add_run(commentary)
        _set_font(comment_run, size=9, italic=True)

        doc.add_paragraph()

    filename = f"Assumption_Setting_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = OUTPUT_DIR / filename
    doc.save(file_path)

    return file_path